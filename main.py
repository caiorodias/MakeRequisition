from docx import Document
from datetime import datetime

# Buscando data

agora = datetime.now()
dia_atual = str(agora.day)
mes_atual = agora.month
if agora.month == 1:
    mes_atual = 'janeiro'
elif agora.month == 2:
    mes_atual = 'fevereiro'
elif agora.month == 3:
    mes_atual = 'março'
elif agora.month == 4:
    mes_atual = 'abril'
elif agora.month == 5:
    mes_atual = 'maio'
elif agora.month == 6:
    mes_atual = 'junho'
elif agora.month == 7:
    mes_atual = 'julho'
elif agora.month == 8:
    mes_atual = 'agosto'
elif agora.month == 9:
    mes_atual = 'setembro'
elif agora.month == 10:
    mes_atual = 'outubro'
elif agora.month == 11:
    mes_atual = 'novembro'
elif agora.month == 12:
    mes_atual = 'dezembro'
else:
    mes_atual == 'DATE_ERROR'
ano_atual = str(agora.year)

# Buscando documento modelo

doc = Document("modelo_rec.docx")

# Reconhecendo tabelas

tabela_preco = doc.tables[0]
tabela_nc = doc.tables[1]

# Apresentação

print('''
Make Requisition

By 3º Sgt Caio Dias
''')

# Edição data

for paragrafo in doc.paragraphs:
    paragrafo.text = paragrafo.text.replace("ediae", dia_atual)
    paragrafo.text = paragrafo.text.replace("emese", mes_atual)
    paragrafo.text = paragrafo.text.replace("eanoe", ano_atual)

# Número da requisição

nr_req = input('Digite o número da requisição:\n')

for paragrafo in doc.paragraphs:
    paragrafo.text = paragrafo.text.replace("ereqe", nr_req)

# Edição assunto

assunto = input('Digite exatamente o assunto da requisição: (sem ponto final)\n')

for paragrafo in doc.paragraphs:
    paragrafo.text = paragrafo.text.replace("eassuntoe", assunto.lower())

# Edição de nomes

aprov = input('Digite o nome completo do APROVISIONADOR:\n')
postoa = input('Digite o posto do APROVISIONADOR (ex. 1º Ten):\n')

fiscal = input('Digite o nome completo do FISCAL ADMINISTRATIVO:\n')
postofa = input('Digite o posto do FISCAL ADMINISTRATIVO (ex. Maj):\n')

od = input('Digite o nome completo do ORDENADOR DE DESPESAS:\n')
postood = input('Digite o posto do ORDENADOR DE DESPESAS (ex. Ten Cel):\n')


for paragrafo in doc.paragraphs:
    
    if '[aprov]' in paragrafo.text:
        texto_aprov = paragrafo.text.replace("[aprov]", aprov.upper())
        paragrafo.clear()
        texto_aprov_negrito = paragrafo.add_run(texto_aprov)
        texto_aprov_negrito.bold = True
    
    if '[postoa]' in paragrafo.text:
        texto_postoa = paragrafo.text.replace("[postoa]", postoa)
        paragrafo.clear()
        texto_postoa_negrito = paragrafo.add_run(texto_postoa)
        texto_postoa_negrito.bold = True
    
    if '[fiscal]' in paragrafo.text:
        texto_fiscal = paragrafo.text.replace("[fiscal]", fiscal.upper())
        paragrafo.clear()
        texto_fiscal_negrito = paragrafo.add_run(texto_fiscal)
        texto_fiscal_negrito.bold = True
    
    if '[postofa]' in paragrafo.text:
        texto_postofa = paragrafo.text.replace("[postofa]", postofa)
        paragrafo.clear()
        texto_postofa_negrito = paragrafo.add_run(texto_postofa)
        texto_postofa_negrito.bold = True
    
    if '[od]' in paragrafo.text:
        texto_od = paragrafo.text.replace("[od]", od.upper())
        paragrafo.clear()
        texto_od_negrito = paragrafo.add_run(texto_od)
        texto_od_negrito.bold = True

    
    if '[postood]' in paragrafo.text:
        texto_postood = paragrafo.text.replace("[postood]", postood)
        paragrafo.clear()
        texto_postood_negrito = paragrafo.add_run(texto_postood)
        texto_postood_negrito.bold = True

# Fornecedor

cnpj = input('Digite o cnpj do fornecedor com todos os caractéres especiais (ex: ., /, -):\n')
fornecedor = input('Digite o nome do fornecedor\n')

for paragrafo in doc.paragraphs:
    paragrafo.text = paragrafo.text.replace("[cnpj]", cnpj)
    paragrafo.text = paragrafo.text.replace("[fornecedor]", fornecedor.upper())

# Subitem

subitem = input('''
Digite o subitem da requisição: 
1 - Combustíveis e Lubrificantes p/ Outras Finalidades
2 - Gás e Outros Materiais Engarrafados
3 - Gêneros de Alimentação
4 - Animais para Pesquisa e Abate
5 - Material de Acondicionamento e Embalagem
6 - Material de Copa e Cozinha
7 - Sementes, Mudas de Plantas e Insumos 
\n''')

if subitem == '1':
    subitem = "03 - COMBUSTÍVEIS E LUBRIFICANTES P/ OUTRAS FINALIDADES"
elif subitem == '2':
    subitem = "04 - GÁS E OUTROS MATERIAIS ENGARRAFADOS"
elif subitem == '3':
    subitem = "07 - GÊNEROS DE ALIMENTAÇÃO"
elif subitem == '4':
    subitem = "08 - ANIMAIS PARA PESQUISA E ABATE"
elif subitem == '5':
    subitem = "19 - MATERIAL DE ACONDICIONAMENTO E EMBALAGEM"
elif subitem == '6':
    subitem = "21 - MATERIAL DE COPA E COZINHA"
elif subitem == '7':
    subitem = "31 - SEMENTES, MUDAS DE PLANTAS E INSUMOS"
else:
    "Porfavor, insira um número válido"

for paragrafo in doc.paragraphs:
    paragrafo.text = paragrafo.text.replace("esie", subitem)

# Nota de Crédito

nc = input('Informe qual nota de crédito será utilizada: (ex: 2023NC403131)\n')
ptres = input('Informe o PTRES da Nota de Crédito: (ex: 193894)\n')
pi = input('Informe qual o PI da Nota de Crédito: (ex: E6SUPJA1QR)\n')

for linha in tabela_nc.rows:
    for celula in linha.cells:
        celula.text = celula.text.replace('eptrese', ptres)
        celula.text = celula.text.replace('ence', nc)
        celula.text = celula.text.replace('epie', pi)

# Salvando

doc.save(f'requisicao_nr{nr_req}.docx')