import docx

doc = docx.Document('modelo_rec.docx')

num_tabelas = len(doc.tables)

for i in range(num_tabelas):
    tabela = doc.tables[i]
    num_linhas = len(tabela.rows)
    num_colunas = len(tabela.columns)
    print(f'Tabela{i}: {num_linhas} linhas x {num_colunas} colunas')